"""Render a Report into JSON, HTML, PDF, and DOCX.

One canonical `Report` (schemas.report) drives every format, so they never
diverge. HTML is the intermediate for PDF (WeasyPrint). DOCX is built directly
with python-docx. PDF rendering is optional at runtime: if WeasyPrint's native
stack is unavailable, `render_pdf` raises a clear error rather than failing
silently.
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from jinja2 import Environment, FileSystemLoader, select_autoescape

from ..core.config import get_settings
from ..core.errors import MetroScanError
from ..schemas.report import Report

_STATUS_LABEL = {
    "compliant": "COMPLIANT",
    "potential_non_compliance": "POTENTIAL NON-COMPLIANCE",
    "not_detected": "NOT DETECTED",
    "not_assessable": "NOT ASSESSABLE",
    "not_applicable": "NOT APPLICABLE",
}


def render_json(report: Report, indent: int = 2) -> str:
    """Canonical machine-readable record."""
    return report.model_dump_json(indent=indent, by_alias=True)


def _environment(template_dir: Optional[Path]) -> Environment:
    template_dir = template_dir or get_settings().report_template_dir
    return Environment(
        loader=FileSystemLoader(str(template_dir)),
        autoescape=select_autoescape(["html", "xml", "j2"]),
    )


def render_html(report: Report, template_dir: Optional[Path] = None) -> str:
    env = _environment(template_dir)
    template = env.get_template("report.html.j2")
    return template.render(r=report)


def render_pdf(report: Report, out_path: Path, template_dir: Optional[Path] = None) -> Path:
    """Render to PDF via WeasyPrint. Raises if WeasyPrint is unavailable."""
    try:
        from weasyprint import HTML  # heavy native deps; import lazily
    except Exception as exc:  # ImportError or native lib load failure
        raise MetroScanError(
            "PDF rendering requires WeasyPrint and its native libraries "
            "(cairo, pango). Install per requirements.txt, or use render_docx/"
            f"render_html. Underlying error: {exc}"
        ) from exc

    html = render_html(report, template_dir)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html).write_pdf(str(out_path))
    return out_path


def render_docx(report: Report, out_path: Path) -> Path:
    """Render an editable DOCX mirroring the report sections."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("MetroScan — Packaged Commodity Compliance Report", level=0)
    doc.add_paragraph(
        f"Ref {report.ref_no or report.report_id} · generated {report.generated_at} · "
        f"app v{report.app_version} · catalog {report.rule_catalog.version}"
    )
    warn = doc.add_paragraph()
    run = warn.add_run(
        "DECISION-SUPPORT — potential non-compliance flagged for officer "
        "verification. Not a final legal finding. Physical verification required "
        "for enforcement."
    )
    run.bold = True

    doc.add_heading("1. Inspection & product", level=1)
    officer = report.inspection.officer
    doc.add_paragraph(f"Officer: {officer.name if officer else '—'}")
    doc.add_paragraph(f"Jurisdiction: {report.inspection.jurisdiction or '—'}")
    doc.add_paragraph(f"Product: {report.product.name or '—'} "
                      f"({report.product.brand or '—'}, {report.product.category or '—'})")

    doc.add_heading("2. Executive summary", level=1)
    s = report.summary
    doc.add_paragraph(
        f"Checked {s.checked} · Compliant {s.compliant} · "
        f"Potential NC {s.potential_non_compliance} · Not detected {s.not_detected} · "
        f"Not assessable {s.not_assessable} · Confidence {s.overall_confidence:.0%}"
    )
    if s.required_actions:
        doc.add_paragraph("Required officer actions:")
        for a in s.required_actions:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("3. Evidence & chain of custody", level=1)
    doc.add_paragraph(f"Original file: {report.evidence.original.file}")
    doc.add_paragraph(f"SHA-256: {report.evidence.original.sha256}")
    doc.add_paragraph(report.evidence.integrity_note)

    doc.add_heading("4. Calibration & measurement basis", level=1)
    c = report.calibration
    doc.add_paragraph(f"Verdict: {c.verdict.value}")
    doc.add_paragraph(f"Reference: {c.reference} "
                      f"({c.aruco_dict} id={c.marker_id})" if c.aruco_dict else
                      f"Reference: {c.reference}")
    if c.mm_per_pixel:
        doc.add_paragraph(f"mm/pixel: {c.mm_per_pixel:.5f} · residual "
                          f"{c.homography_residual_px:.2f}px")
    if c.verdict.value != "calibrated":
        doc.add_paragraph(f"Uncalibrated: {c.reason} — no millimetre verdicts reported.")

    doc.add_heading("5. Declaration findings (Rule 6)", level=1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Declaration", "Clause", "Extracted", "Status", "Note"]):
        hdr[i].text = h
    for d in report.declarations:
        row = t.add_row().cells
        row[0].text = d.label
        row[1].text = d.clause_ref.clause
        row[2].text = d.extracted or "—"
        row[3].text = _STATUS_LABEL.get(d.status.value, d.status.value)
        row[4].text = d.note or ""

    doc.add_heading("6. Font-size & placement analysis (Rule 7)", level=1)
    fa = report.font_analysis
    if fa.panel_area_cm2 and fa.table_i_band:
        doc.add_paragraph(
            f"Panel area: {fa.panel_area_cm2.value:.1f} ± {fa.panel_area_cm2.uncertainty:.1f} cm² "
            f"→ band {fa.table_i_band.area_band}, min height {fa.table_i_band.min_height_mm} mm"
        )
    if fa.items:
        ft = doc.add_table(rows=1, cols=5)
        ft.style = "Light Grid Accent 1"
        fh = ft.rows[0].cells
        for i, h in enumerate(["Declaration", "Height", "Threshold", "Status", "Reason"]):
            fh[i].text = h
        for it in fa.items:
            row = ft.add_row().cells
            row[0].text = it.declaration_id
            row[1].text = (f"{it.height_mm.value:.2f} ± {it.height_mm.uncertainty:.2f} mm"
                           if it.height_mm else "—")
            row[2].text = f"{it.threshold_mm:.1f} mm" if it.threshold_mm is not None else "—"
            row[3].text = _STATUS_LABEL.get(it.status.value, it.status.value)
            row[4].text = it.reason or ""

    doc.add_heading("7. Limitations & confidence", level=1)
    doc.add_paragraph(report.limitations)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
